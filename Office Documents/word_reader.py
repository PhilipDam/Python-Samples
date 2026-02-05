#!/usr/bin/env python3
"""
================================================================================
docx_format_breaks.commentary.py
================================================================================

GOAL
----
Take a Microsoft Word .docx file as input and emit a *single text stream* as output,
but insert explicit START/END markers whenever the formatting context changes.

This is useful when you want to rebuild the Word document into your own intermediate
representation (IR), Markdown-ish format, or any downstream parser that needs to
know exactly where formatting begins/ends.

WHAT THIS SCRIPT EMITS
----------------------
The output is plain text that contains:

1) Document boundaries
   [[DOC_START ...]]
   [[DOC_END]]

2) Paragraph boundaries + paragraph "context"
   [[P_START ...]]
   paragraph text...
   [[P_END]]

   The P_START marker includes:
     - paragraph style name
     - paragraph justification (alignment): left/center/right/justify/...
     - list info, if paragraph is in a list:
         list kind (bullet/numbered/list)
         list nesting level (ilvl)
         Word numbering format (numFmt), e.g. "decimal", "lowerRoman", ...
         For numbered lists: the actual item number emitted: itemNo=<n>

3) Run-level markers (within paragraph text)
   Bold / Italic are emitted as explicit nested markers:
     [[B_START]] ... [[B_END]]
     [[I_START]] ... [[I_END]]

   Run style (character style) is emitted as:
     [[RSTYLE_START name="..."]] ... [[RSTYLE_END name="..."]]

4) Lists (start/end markers separate from paragraph markers)
   When the traversal enters a new list instance (Word numId changes):
     [[LIST_START kind="numbered" numId="5" start=3 continued="true"]]
   When leaving a list:
     [[LIST_END kind="numbered" numId="5"]]

   Notes:
   - "start" is the configured start number for that list instance and level,
     derived from startOverride / start.
   - "continued" is a *heuristic* (best-effort): if we have seen the same numId earlier
     in the document traversal, we mark continued="true". This is a good signal that
     Word may be continuing a previous list after some intervening content.

5) Tables
   We mark the table/row/cell boundaries:
     [[TABLE_START]]
       [[ROW_START r=0]]
         [[CELL_START r=0 c=0]]
           ... paragraphs/tables within the cell ...
         [[CELL_END r=0 c=0]]
       [[ROW_END r=0]]
     [[TABLE_END]]

IMPORTANT REALITY CHECK
-----------------------
Word .docx formatting is not "simple tags". It's a web of:
- direct formatting on runs/paragraphs,
- styles (character + paragraph),
- style inheritance ("basedOn"),
- numbering definitions split across multiple XML parts,
- overrides, continuations, multi-level list patterns, etc.

This script aims to be:
- deterministic
- traceable
- best-effort faithful
without trying to replicate *every* nuance of Word's layout engine.

In particular:
- "Effective" bold/italic resolution is approximated (run -> run style -> paragraph style -> Normal).
- Actual list *labels* (like "IV.", "a)", "1.2.3") are not fully rendered;
  instead we output the *actual numeric counter* (itemNo) and the raw numFmt.

DEPENDENCIES
------------
pip install python-docx

USAGE
-----
python word_reader.py input.docx > out.txt
python word_reader.py input.docx -o out.txt

================================================================================
"""

from __future__ import annotations

import argparse
import sys
from dataclasses import dataclass
from typing import Dict, Iterator, List, Optional, Tuple, Union, Set

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

# python-docx is a high-level library, but some features (like numbering and
# justification) are not fully exposed as convenient properties.
# We therefore access the underlying WordprocessingML XML via ._element / ._p
# and use qn(...) to build the fully qualified names ("w:p", "w:tbl", etc).
from docx.oxml.ns import qn


# =============================================================================
# SECTION 1: Output helpers
# =============================================================================

def emit(out, s: str) -> None:
    """
    Write a string to the output stream (no newline added).

    Keeping this as a helper function makes it easy to:
    - redirect output to files,
    - implement buffering,
    - insert logging,
    - or change output format later.
    """
    out.write(s)


def emit_line(out, s: str) -> None:
    """
    Write a line to the output stream, always ending with newline.
    """
    out.write(s + "\n")


def esc_attr(s: str) -> str:
    """
    Escape a string so it can safely appear in marker attributes.

    We keep escaping minimal and human-readable:
    - only escape quotes to avoid breaking attribute syntax.

    Example:
        style name:  My "Special" Style
        becomes:     My \\"Special\\" Style
    """
    return (s or "").replace('"', '\\"')


# =============================================================================
# SECTION 2: Iterating document blocks in "visual" order
# =============================================================================
#
# Word documents are not stored as a simple list of paragraphs.
# The document body contains a sequence of blocks, which can be:
#   - paragraphs (<w:p>)
#   - tables     (<w:tbl>)
#
# Each table contains rows/cells; each cell can contain paragraphs AND tables
# (nested tables).
#
# python-docx exposes:
#   doc.paragraphs
#   doc.tables
# but those lists do *not* preserve the interleaving order.
#
# Therefore, we iterate the underlying XML children of the document body (or cell)
# and yield Paragraph/Table objects in the exact order they appear in the XML.
# =============================================================================

Block = Union[Paragraph, Table]


def iter_block_items(parent: Union[DocxDocument, _Cell]) -> Iterator[Block]:
    """
    Yield Paragraph and Table objects in the order they appear in the given parent.

    parent can be:
      - the whole Document
      - a table Cell

    Implementation details:
    - For documents: parent.element.body contains the <w:body>.
    - For cells: parent._tc contains the <w:tc> element.
    - We iterate children and wrap them into python-docx objects.

    This ordering is crucial for your use case because you want a single stream
    of content that matches the reading order.
    """
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc

    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            # Wrap raw <w:p> into a Paragraph object.
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            # Wrap raw <w:tbl> into a Table object.
            yield Table(child, parent)


# =============================================================================
# SECTION 3: Numbering / list model (the infamous "numbering.xml")
# =============================================================================
#
# Word lists are not stored as “this paragraph is bullet” in a single attribute.
#
# Instead, a list paragraph typically has:
#   <w:pPr>
#     <w:numPr>
#       <w:numId w:val="5"/>
#       <w:ilvl  w:val="0"/>
#     </w:numPr>
#   </w:pPr>
#
# Where:
# - numId identifies a *numbering instance* (<w:num>).
# - ilvl is the nesting level (0 = top, 1 = nested, etc.)
#
# That numId points to an abstract definition (<w:abstractNum>),
# which describes the formatting per level:
#   - numFmt: bullet / decimal / roman / etc.
#   - start:  initial number for that level
#   - lvlText: pattern (e.g., "%1.") -- NOT rendered fully in this script
#
# On top of that, a <w:num> instance can override the start per level:
#   <w:lvlOverride w:ilvl="0">
#      <w:startOverride w:val="5"/>
#   </w:lvlOverride>
#
# For your requirement:
# - list kind (bullet/numbered)
# - configured start number
# - actual item numbers per paragraph
#
# we need:
#   numId -> abstractNumId
#   abstractNumId -> (ilvl -> numFmt)
#   abstractNumId -> (ilvl -> start)
#   numId -> (ilvl -> startOverride)
# =============================================================================

@dataclass(frozen=True)
class ListInfo:
    """
    Per-paragraph list metadata extracted from <w:numPr> + numbering model.

    num_id:
      Word numbering instance ID (stringified)
    ilvl:
      nesting level, 0..8 typically
    fmt:
      Word numFmt value such as "bullet", "decimal", "lowerRoman", ...
    kind:
      "bullet" | "numbered" | "list"
      We classify fmt into a simpler category.
    """
    num_id: str
    ilvl: int
    fmt: str
    kind: str


def _safe_int(x: Optional[str], default: int = 0) -> int:
    """
    Convert a string to int safely, returning default on None/bad input.
    """
    try:
        return int(x) if x is not None else default
    except ValueError:
        return default


@dataclass
class NumberingModel:
    """
    A condensed, query-friendly view of numbering.xml.

    We store just enough to:
      - determine list type (bullet vs numbered) per numId+level
      - determine start numbers (overrides + defaults)
    """
    numId_to_abstractId: Dict[str, str]
    abstractId_to_lvlFmt: Dict[str, Dict[int, str]]
    abstractId_to_lvlStart: Dict[str, Dict[int, int]]
    numId_to_lvlStartOverride: Dict[str, Dict[int, int]]


def build_numbering_model(doc: Document) -> NumberingModel:
    """
    Parse the document's numbering part (if present) and build the NumberingModel.

    If the document has no numbering part, we return empty maps and list detection
    will simply yield "not a list" for all paragraphs.
    """
    numId_to_abstractId: Dict[str, str] = {}
    abstractId_to_lvlFmt: Dict[str, Dict[int, str]] = {}
    abstractId_to_lvlStart: Dict[str, Dict[int, int]] = {}
    numId_to_lvlStartOverride: Dict[str, Dict[int, int]] = {}

    # Not all documents have numbering (no lists => no numbering part).
    try:
        numbering_part = doc.part.numbering_part
    except Exception:
        return NumberingModel(
            numId_to_abstractId, abstractId_to_lvlFmt, abstractId_to_lvlStart, numId_to_lvlStartOverride
        )

    root = numbering_part._element  # underlying XML root

    # -------------------------------------------------------------------------
    # 1) Parse <w:num> instances:
    #    <w:num w:numId="5">
    #      <w:abstractNumId w:val="3"/>
    #      <w:lvlOverride w:ilvl="0">
    #        <w:startOverride w:val="5"/>
    #      </w:lvlOverride>
    #    </w:num>
    # -------------------------------------------------------------------------
    for num in root.findall(qn("w:num")):
        num_id = num.get(qn("w:numId"))
        if not num_id:
            continue

        abs_elm = num.find(qn("w:abstractNumId"))
        abs_id = abs_elm.get(qn("w:val")) if abs_elm is not None else None
        if abs_id:
            numId_to_abstractId[num_id] = abs_id

        # startOverride per level, if present
        for lvl_override in num.findall(qn("w:lvlOverride")):
            ilvl = _safe_int(lvl_override.get(qn("w:ilvl")), 0)
            start_ovr = lvl_override.find(qn("w:startOverride"))
            if start_ovr is not None:
                val = start_ovr.get(qn("w:val"))
                if val is not None:
                    numId_to_lvlStartOverride.setdefault(num_id, {})[ilvl] = _safe_int(val, 1)

    # -------------------------------------------------------------------------
    # 2) Parse <w:abstractNum> definitions:
    #    <w:abstractNum w:abstractNumId="3">
    #      <w:lvl w:ilvl="0">
    #         <w:numFmt w:val="decimal"/>
    #         <w:start  w:val="1"/>
    #      </w:lvl>
    #      ...
    #    </w:abstractNum>
    # -------------------------------------------------------------------------
    for absnum in root.findall(qn("w:abstractNum")):
        abs_id = absnum.get(qn("w:abstractNumId"))
        if not abs_id:
            continue

        lvl_fmt: Dict[int, str] = {}
        lvl_start: Dict[int, int] = {}

        for lvl in absnum.findall(qn("w:lvl")):
            ilvl = _safe_int(lvl.get(qn("w:ilvl")), 0)

            fmt_elm = lvl.find(qn("w:numFmt"))
            fmt = fmt_elm.get(qn("w:val")) if fmt_elm is not None else ""
            if fmt:
                lvl_fmt[ilvl] = fmt

            start_elm = lvl.find(qn("w:start"))
            if start_elm is not None:
                sval = start_elm.get(qn("w:val"))
                if sval is not None:
                    lvl_start[ilvl] = _safe_int(sval, 1)

        if lvl_fmt:
            abstractId_to_lvlFmt[abs_id] = lvl_fmt
        if lvl_start:
            abstractId_to_lvlStart[abs_id] = lvl_start

    return NumberingModel(
        numId_to_abstractId=numId_to_abstractId,
        abstractId_to_lvlFmt=abstractId_to_lvlFmt,
        abstractId_to_lvlStart=abstractId_to_lvlStart,
        numId_to_lvlStartOverride=numId_to_lvlStartOverride,
    )


def paragraph_list_info(p: Paragraph, nm: NumberingModel) -> Optional[ListInfo]:
    """
    Determine whether a paragraph is part of a list.

    If yes, return a ListInfo describing the list instance (numId) and nesting level (ilvl),
    plus format classification derived from the numbering model.

    If no, return None.
    """
    # Underlying XML paragraph properties:
    pPr = p._p.pPr
    if pPr is None or pPr.numPr is None:
        return None

    # Grab numId and ilvl from <w:numPr>
    numId_elm = pPr.numPr.numId
    ilvl_elm = pPr.numPr.ilvl

    num_id = numId_elm.val if numId_elm is not None else None
    ilvl = int(ilvl_elm.val) if ilvl_elm is not None and ilvl_elm.val is not None else 0
    if not num_id:
        return None

    # Determine numFmt (bullet/decimal/etc.) from numId -> abstractNum -> lvl
    abs_id = nm.numId_to_abstractId.get(str(num_id))
    fmt = ""
    if abs_id:
        fmt = nm.abstractId_to_lvlFmt.get(abs_id, {}).get(ilvl, "")

    # Reduce Word's many numFmt values to a friendlier category
    if fmt == "bullet":
        kind = "bullet"
    elif fmt in (
        "decimal", "decimalZero",
        "upperRoman", "lowerRoman",
        "upperLetter", "lowerLetter",
        "ordinal", "cardinalText",
        "ordinalText", "hex",
        "chineseCounting", "aiueo", "iroha",
    ):
        kind = "numbered"
    else:
        # Could still be numbered with some exotic format, but we label as generic "list".
        kind = "list"

    return ListInfo(num_id=str(num_id), ilvl=ilvl, fmt=fmt or "unknown", kind=kind)


def list_configured_start(num_id: str, ilvl: int, nm: NumberingModel) -> int:
    """
    Get the configured *start number* for a list instance and level.

    Priority order matches Word semantics:
    1) A per-instance override (<w:startOverride>) on <w:num>/<w:lvlOverride>
    2) The abstract definition default (<w:start>) on <w:abstractNum>/<w:lvl>
    3) Otherwise default to 1
    """
    ov = nm.numId_to_lvlStartOverride.get(num_id, {}).get(ilvl)
    if ov is not None:
        return int(ov)

    abs_id = nm.numId_to_abstractId.get(num_id)
    if abs_id:
        st = nm.abstractId_to_lvlStart.get(abs_id, {}).get(ilvl)
        if st is not None:
            return int(st)

    return 1


# =============================================================================
# SECTION 4: Paragraph justification (alignment)
# =============================================================================
#
# Word stores paragraph alignment in:
#   <w:pPr>
#     <w:jc w:val="center"/>
#   </w:pPr>
#
# It can also be specified in the paragraph style.
#
# python-docx has paragraph.alignment, but it may not always reflect style inheritance
# the way you want, and it can be an enum. We read the XML directly and do a
# best-effort resolution:
#   direct paragraph jc -> paragraph style jc -> Normal style jc -> "left"
# =============================================================================

def _style_name(style) -> str:
    """
    Return the display name of a python-docx style object, or "" if unavailable.
    """
    try:
        return style.name or ""
    except Exception:
        return ""

def _style_id_from_style_obj(style) -> str:
    """Return style.style_id, or ""."""
    try:
        return getattr(style, "style_id", "") or ""
    except Exception:
        return ""


def paragraph_style_name(p: Paragraph, doc: Document) -> str:
    """Robustly resolve a paragraph's *display* style name."""
    try:
        if getattr(p, "style", None) is not None and p.style.name:
            return p.style.name
    except Exception:
        pass
    try:
        pPr = p._p.pPr
        if pPr is not None and pPr.pStyle is not None and pPr.pStyle.val is not None:
            sid = str(pPr.pStyle.val)
            try:
                return doc.styles[sid].name or sid
            except Exception:
                return sid
    except Exception:
        pass
    return ""


def run_character_style_name(run, doc: Document) -> str:
    """Robustly resolve a run's character style name (if any)."""
    try:
        rs = getattr(run, "style", None)
        if rs is not None and rs.name:
            return rs.name
    except Exception:
        pass
    try:
        rPr = run._r.rPr
        if rPr is not None:
            rStyle = rPr.find(qn("w:rStyle"))
            if rStyle is not None:
                sid = rStyle.get(qn("w:val")) or ""
                if sid:
                    try:
                        return doc.styles[sid].name or sid
                    except Exception:
                        return sid
    except Exception:
        pass
    return ""


def _norm_jc(val: Optional[str]) -> str:
    """
    Normalize Word's jc values to a stable vocabulary.

    Common Word values:
      left, center, right, both (justified), distribute, start, end

    We normalize "both" -> "justify" to reduce surprises.
    """
    if not val:
        return "inherit"
    v = val.lower()
    mapping = {
        "left": "left",
        "center": "center",
        "right": "right",
        "both": "justify",
        "justify": "justify",
        "distribute": "distribute",
        "start": "start",
        "end": "end",
    }
    return mapping.get(v, v)


def _ppr_jc_value(p: Paragraph) -> Optional[str]:
    """
    Read <w:jc w:val="..."> from this specific paragraph's properties (direct formatting).
    """
    pPr = p._p.pPr
    if pPr is None:
        return None
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        return None
    return jc.get(qn("w:val"))


def effective_paragraph_justification(p: Paragraph, doc: Document) -> str:
    """
    Best-effort resolution for paragraph justification:
      1) direct pPr jc
      2) paragraph style pPr jc
      3) Normal style pPr jc
      4) default "left"
    """
    v = _ppr_jc_value(p)
    if v:
        return _norm_jc(v)

    def style_jc(style_obj) -> Optional[str]:
        """
        Read jc from a style's <w:pPr>.
        """
        try:
            elm = style_obj._element
        except Exception:
            return None
        pPr = elm.find(qn("w:pPr"))
        if pPr is None:
            return None
        jc = pPr.find(qn("w:jc"))
        if jc is None:
            return None
        return jc.get(qn("w:val"))

    para_style = getattr(p, "style", None)
    if para_style is not None:
        v2 = style_jc(para_style)
        if v2:
            return _norm_jc(v2)

    try:
        normal_style = doc.styles["Normal"]
        v3 = style_jc(normal_style)
        if v3:
            return _norm_jc(v3)
    except Exception:
        pass

    return "left"


# =============================================================================
# SECTION 5: Run formatting (bold/italic + run character style)
# =============================================================================
#
# A paragraph is broken into "runs". Each run can have its own formatting.
#
# python-docx gives run.bold and run.italic, but those can be None to indicate
# "inherit". Effective formatting can be inherited from:
#   run -> run style -> paragraph style -> Normal
#
# We implement a modest inheritance chain to get deterministic boolean values.
# =============================================================================

@dataclass(frozen=True)
class RunFormat:
    """
    Computed (effective-ish) run formatting state.

    bold/italic:
      Concrete booleans after resolving inheritance (best effort).

    rstyle:
      Character style name applied to this run ("" if none).
    """
    bold: bool
    italic: bool
    rstyle: str


def resolve_bool_chain(*vals: Optional[bool], default: bool = False) -> bool:
    """
    Given values that might be None (meaning "inherit"),
    return the first non-None value. If all are None, return default.
    """
    for v in vals:
        if v is not None:
            return bool(v)
    return default


def effective_run_format(run, para: Paragraph, doc: Document) -> RunFormat:
    """
    Best-effort run format resolution:
      run.bold/italic ->
      run.style.font.bold/italic ->
      para.style.font.bold/italic ->
      Normal.style.font.bold/italic ->
      False

    This will not cover *every* Word inheritance scenario, but it covers most
    real-world documents well enough for a formatting-break stream.
    """
    run_style = getattr(run, "style", None)
    para_style = getattr(para, "style", None)
    normal_style = None
    try:
        normal_style = doc.styles["Normal"]
    except Exception:
        normal_style = None

    bold = resolve_bool_chain(
        run.bold,
        getattr(getattr(run_style, "font", None), "bold", None),
        getattr(getattr(para_style, "font", None), "bold", None),
        getattr(getattr(normal_style, "font", None), "bold", None),
        default=False,
    )
    italic = resolve_bool_chain(
        run.italic,
        getattr(getattr(run_style, "font", None), "italic", None),
        getattr(getattr(para_style, "font", None), "italic", None),
        getattr(getattr(normal_style, "font", None), "italic", None),
        default=False,
    )

    rstyle_name = run_character_style_name(run, doc)
    # Word often assigns 'Default Paragraph Font' as the implicit run style.
    # It is not helpful as a semantic marker, so suppress it.
    if rstyle_name == "Default Paragraph Font":
        rstyle_name = ""
    return RunFormat(bold=bold, italic=italic, rstyle=rstyle_name)


# =============================================================================
# SECTION 6: Streaming state + marker emission logic
# =============================================================================
#
# The core problem is: "insert a break each time formatting changes".
# That is a streaming/differencing problem.
#
# We maintain a StreamState representing the currently-open markers
# (bold, italic, run style) and the current paragraph/list/table context.
#
# Whenever we move to a new paragraph or a new run, we compare the new desired
# format to the current state and emit END/START markers as necessary.
# =============================================================================

@dataclass
class StreamState:
    """
    The full "current context" while traversing the document.

    - in_table: nesting depth for tables (informational; not heavily used)
    - pstyle/pjc: paragraph style name and justification
    - in_list/list_*: for emitting LIST_START/LIST_END markers
    - seen_num_ids: for continued list heuristic
    - counters: numbering counters for actual item numbers
    - bold/italic/rstyle: currently-open run markers
    """
    in_table: int = 0

    # Paragraph context
    pstyle: str = ""
    pjc: str = ""

    # List context (for markers)
    in_list: bool = False
    list_kind: str = ""
    list_num_id: str = ""
    list_lvl: int = 0

    # List tracking
    seen_num_ids: Set[str] = None
    counters: Dict[Tuple[str, int], int] = None  # (numId, ilvl) -> next number to emit

    # Run context
    bold: bool = False
    italic: bool = False
    rstyle: str = ""


def close_run_markers(out, st: StreamState) -> None:
    """
    Close any currently-open run-level markers in a deterministic order.

    We close:
      run style -> italic -> bold

    The order isn't mandated by Word, but it's tidy and makes the output easy to parse.
    """
    if st.rstyle:
        emit(out, f'[[RSTYLE_END name="{esc_attr(st.rstyle)}"]]')
        st.rstyle = ""
    if st.italic:
        emit(out, "[[I_END]]")
        st.italic = False
    if st.bold:
        emit(out, "[[B_END]]")
        st.bold = False


def switch_run_format(out, st: StreamState, new_fmt: RunFormat) -> None:
    """
    Transition from current run formatting (st) to the desired new formatting (new_fmt)
    by emitting END/START markers only where needed.

    This is the heart of the "break when formatting changes" requirement.
    """
    # Handle character style change first (close old if different).
    if st.rstyle != new_fmt.rstyle:
        if st.rstyle:
            emit(out, f'[[RSTYLE_END name="{esc_attr(st.rstyle)}"]]')
        st.rstyle = ""

    # Close markers that should end
    if st.italic and not new_fmt.italic:
        emit(out, "[[I_END]]")
        st.italic = False
    if st.bold and not new_fmt.bold:
        emit(out, "[[B_END]]")
        st.bold = False

    # Open markers that should begin
    if new_fmt.bold and not st.bold:
        emit(out, "[[B_START]]")
        st.bold = True
    if new_fmt.italic and not st.italic:
        emit(out, "[[I_START]]")
        st.italic = True

    # Open new character style if needed
    if new_fmt.rstyle and new_fmt.rstyle != st.rstyle:
        emit(out, f'[[RSTYLE_START name="{esc_attr(new_fmt.rstyle)}"]]')
        st.rstyle = new_fmt.rstyle


def emit_list_start(out, li: ListInfo, start_num: int, continued: bool) -> None:
    """
    Emit a LIST_START marker when we enter a new list instance (numId changes).

    start_num:
      Configured start for this list instance and level (from startOverride or start).
    continued:
      Heuristic: have we seen this numId earlier in traversal?
    """
    emit_line(
        out,
        f'[[LIST_START kind="{li.kind}" numId="{li.num_id}" start={start_num} continued="{str(continued).lower()}"]]'
    )


def emit_list_end(out, st: StreamState) -> None:
    """
    Emit LIST_END marker for the currently active list instance.
    """
    emit_line(out, f'[[LIST_END kind="{st.list_kind}" numId="{st.list_num_id}"]]')


def ensure_counter_initialized(st: StreamState, num_id: str, ilvl: int, nm: NumberingModel) -> None:
    """
    Ensure st.counters contains an entry for (num_id, ilvl).
    If missing, initialize it to the configured start number.
    """
    key = (num_id, ilvl)
    if key not in st.counters:
        st.counters[key] = list_configured_start(num_id, ilvl, nm)


def reset_deeper_levels(st: StreamState, num_id: str, parent_lvl: int, nm: NumberingModel) -> None:
    """
    Common nested-list behavior:
    When a parent level item increments, deeper levels usually restart.

    Example (outline numbering):
      1.
         a.
         b.
      2.       <-- deeper levels typically restart for the new parent item
         a.

    We implement this by resetting counters for all (num_id, lvl) where lvl > parent_lvl
    back to their configured start.
    """
    to_reset = [lvl for (nid, lvl) in st.counters.keys() if nid == num_id and lvl > parent_lvl]
    for lvl in to_reset:
        st.counters[(num_id, lvl)] = list_configured_start(num_id, lvl, nm)


def next_item_number(st: StreamState, li: ListInfo, nm: NumberingModel) -> Optional[int]:
    """
    Compute the *actual* item number for this paragraph IF it's a numbered list item.

    - If not numbered: return None.
    - Else:
        - read current counter
        - emit that as item number
        - increment counter for next time
        - reset deeper levels

    This produces a robust approximation of "what number would Word show"
    for the numeric portion of the label.
    """
    if li.kind != "numbered":
        return None

    ensure_counter_initialized(st, li.num_id, li.ilvl, nm)

    n = st.counters[(li.num_id, li.ilvl)]
    st.counters[(li.num_id, li.ilvl)] = n + 1

    reset_deeper_levels(st, li.num_id, li.ilvl, nm)

    return n


def start_paragraph(out, st: StreamState, pstyle: str, pjc: str, li: Optional[ListInfo], nm: NumberingModel) -> None:
    """
    Enter a paragraph:
    - close any run markers still open from previous paragraph
    - optionally emit LIST_END / LIST_START if list context changes
    - compute list item number if needed
    - emit P_START marker carrying paragraph context
    """
    close_run_markers(out, st)
    emit(out, "\n")  # paragraph boundary as a blank line separator (improves readability)

    new_in_list = li is not None

    # If we were in a list and now are not, close the list marker.
    if st.in_list and not new_in_list:
        emit_list_end(out, st)
        st.in_list = False
        st.list_kind = ""
        st.list_num_id = ""
        st.list_lvl = 0

    # If we are in a list now, possibly start a new list marker if numId changes.
    if new_in_list and li is not None:
        need_new_list_marker = (not st.in_list) or (st.list_num_id != li.num_id)

        if need_new_list_marker:
            continued = li.num_id in st.seen_num_ids
            st.seen_num_ids.add(li.num_id)

            start_num = list_configured_start(li.num_id, li.ilvl, nm)
            emit_list_start(out, li, start_num=start_num, continued=continued)

            st.in_list = True
            st.list_kind = li.kind
            st.list_num_id = li.num_id
            st.list_lvl = li.ilvl
        else:
            # Same list instance, but level might have changed.
            st.list_lvl = li.ilvl

    # Store paragraph state
    st.pstyle = pstyle or ""
    st.pjc = pjc or "inherit"

    # Compute per-item number if this is a numbered list paragraph.
    item_no = next_item_number(st, li, nm) if li is not None else None

    # Emit P_START marker (one marker per paragraph, carrying all paragraph-level data)
    if li is not None:
        if item_no is not None:
            emit_line(
                out,
                f'[[P_START style="{esc_attr(st.pstyle)}" jc="{esc_attr(st.pjc)}" '
                f'list="{li.kind}" lvl={li.ilvl} fmt="{esc_attr(li.fmt)}" itemNo={item_no}]]'
            )
        else:
            emit_line(
                out,
                f'[[P_START style="{esc_attr(st.pstyle)}" jc="{esc_attr(st.pjc)}" '
                f'list="{li.kind}" lvl={li.ilvl} fmt="{esc_attr(li.fmt)}"]]'
            )
    else:
        emit_line(out, f'[[P_START style="{esc_attr(st.pstyle)}" jc="{esc_attr(st.pjc)}"]]')


def end_paragraph(out, st: StreamState) -> None:
    """
    End a paragraph:
    - close any run markers that are still open
    - emit P_END marker
    """
    close_run_markers(out, st)
    emit_line(out, "[[P_END]]")


def emit_paragraph_text(out, st: StreamState, doc: Document, p: Paragraph) -> None:
    """
    Emit the actual textual content of a paragraph.

    We stream over runs in order. For each run:
      - compute effective formatting
      - switch marker state as needed
      - emit the run text

    This guarantees that formatting changes are reflected by explicit markers.
    """
    for run in p.runs:
        txt = run.text
        if not txt:
            continue

        new_fmt = effective_run_format(run, p, doc)
        switch_run_format(out, st, new_fmt)
        emit(out, txt)


def emit_table(out, st: StreamState, doc: Document, table: Table, nm: NumberingModel) -> None:
    """
    Emit a table structure including:
      - TABLE_START / TABLE_END
      - ROW_START / ROW_END
      - CELL_START / CELL_END

    And recursively emit the content inside each cell (paragraphs + nested tables).
    """
    close_run_markers(out, st)
    emit_line(out, "[[TABLE_START]]")
    st.in_table += 1

    for r_i, row in enumerate(table.rows):
        emit_line(out, f"[[ROW_START r={r_i}]]")
        for c_i, cell in enumerate(row.cells):
            emit_line(out, f"[[CELL_START r={r_i} c={c_i}]]")

            # A cell can contain multiple paragraphs and nested tables.
            for block in iter_block_items(cell):
                if isinstance(block, Paragraph):
                    pstyle = paragraph_style_name(block, doc)
                    pjc = effective_paragraph_justification(block, doc)
                    li = paragraph_list_info(block, nm)

                    start_paragraph(out, st, pstyle, pjc, li, nm)
                    emit_paragraph_text(out, st, doc, block)
                    end_paragraph(out, st)
                else:
                    # Nested table inside a cell.
                    emit_table(out, st, doc, block, nm)

            emit_line(out, f"[[CELL_END r={r_i} c={c_i}]]")
        emit_line(out, f"[[ROW_END r={r_i}]]")

    st.in_table -= 1
    emit_line(out, "[[TABLE_END]]")


# =============================================================================
# SECTION 7: High-level document traversal
# =============================================================================

def process_document(docx_path: str, out) -> None:
    """
    Main processing function:
      - load the .docx
      - build numbering model
      - traverse blocks in order
      - emit markers + text
    """
    doc = Document(docx_path)
    nm = build_numbering_model(doc)

    # Initialize stream state with empty list-tracking and counters.
    st = StreamState(seen_num_ids=set(), counters={})

    emit_line(out, f'[[DOC_START path="{esc_attr(docx_path)}"]]')
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            pstyle = paragraph_style_name(block, doc)
            pjc = effective_paragraph_justification(block, doc)
            li = paragraph_list_info(block, nm)

            start_paragraph(out, st, pstyle, pjc, li, nm)
            emit_paragraph_text(out, st, doc, block)
            end_paragraph(out, st)
        else:
            emit_table(out, st, doc, block, nm)

    # If the document ends while we're inside a list marker, close it.
    if st.in_list:
        emit_list_end(out, st)
        st.in_list = False

    # Close any lingering run markers (paranoia / correctness)
    close_run_markers(out, st)
    emit_line(out, "[[DOC_END]]")


# =============================================================================
# SECTION 8: CLI entrypoint
# =============================================================================

def main(argv: List[str]) -> int:
    """
    Command line interface.

    We keep the script CLI-friendly:
      docx_format_breaks.commentary.py input.docx
      docx_format_breaks.commentary.py input.docx -o out.txt
    """
    ap = argparse.ArgumentParser(
        description="Emit Word .docx text with explicit formatting markers (incl. justification + actual list item numbers)."
    )
    ap.add_argument("docx", help="Input .docx file path")
    ap.add_argument("-o", "--output", default="", help="Optional output file path (defaults to stdout).")
    args = ap.parse_args(argv)

    if args.output:
        with open(args.output, "w", encoding="utf-8") as f:
            process_document(args.docx, f)
    else:
        process_document(args.docx, sys.stdout)

    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))